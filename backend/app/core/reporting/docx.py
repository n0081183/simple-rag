"""DOCX report export with tables and formatting."""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt

from app.core.reporting.anonymizer import anonymize_results
from app.core.verification.schemas import VerificationResult, Verdict

_LABELS = {
    "pl": {
        "title": "Raport weryfikacji wymagań SIWZ",
        "summary": "Podsumowanie",
        "requirement": "Wymaganie",
        "verdict": "Ocena",
        "confidence": "Pewność",
        "reasoning": "Uzasadnienie",
        "evidence": "Dowody",
        "source": "Źródło",
        "quote": "Cytat",
        "auto_warn": "Produkt dopasowany automatycznie — wymaga weryfikacji przez specjalistę Cortex.",
        "met": "Spełnione",
        "partial": "Częściowo",
        "not_met": "Niespełnione",
        "unclear": "Niejasne",
    },
    "en": {
        "title": "SIWZ Requirements Verification Report",
        "summary": "Summary",
        "requirement": "Requirement",
        "verdict": "Verdict",
        "confidence": "Confidence",
        "reasoning": "Reasoning",
        "evidence": "Evidence",
        "source": "Source",
        "quote": "Quote",
        "auto_warn": "Product was auto-detected — requires expert verification by a Cortex specialist.",
        "met": "Met",
        "partial": "Partial",
        "not_met": "Not met",
        "unclear": "Unclear",
    },
}

_VERDICT_LABEL = {
    Verdict.MET: "met",
    Verdict.PARTIAL: "partial",
    Verdict.NOT_MET: "not_met",
    Verdict.UNCLEAR: "unclear",
}


def export_docx(
    results: list[VerificationResult],
    title: str | None = None,
    language: str = "pl",
    auto_detect_warning: bool = False,
    product: str | None = None,
    anonymize: bool = False,
) -> bytes:
    L = _LABELS.get(language, _LABELS["en"])
    if anonymize:
        results = anonymize_results(results)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    doc.add_heading(title or L["title"], 0)
    if product:
        doc.add_paragraph(f"Product: {product}")
    if auto_detect_warning:
        p = doc.add_paragraph(L["auto_warn"])
        p.runs[0].bold = True

    # Summary table
    counts = {v: 0 for v in Verdict}
    for r in results:
        counts[r.verdict] += 1
    doc.add_heading(L["summary"], level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = L["verdict"]
    hdr[1].text = "Count"
    for verdict in Verdict:
        row = table.add_row().cells
        key = _VERDICT_LABEL[verdict]
        row[0].text = L.get(key, verdict.value)
        row[1].text = str(counts[verdict])

    for i, r in enumerate(results, 1):
        doc.add_heading(f"{i}. {r.verdict.value} — {L['confidence']}: {r.confidence.value}", level=2)
        doc.add_paragraph(r.requirement_text)
        if r.reasoning_steps:
            doc.add_paragraph(L["reasoning"], style="Heading 3")
            for step in r.reasoning_steps:
                doc.add_paragraph(step, style="List Bullet")
        if r.evidence:
            doc.add_paragraph(L["evidence"], style="Heading 3")
            ev_table = doc.add_table(rows=1, cols=2)
            ev_table.style = "Table Grid"
            ev_table.rows[0].cells[0].text = L["source"]
            ev_table.rows[0].cells[1].text = L["quote"]
            for ev in r.evidence:
                cells = ev_table.add_row().cells
                cells[0].text = ev.source_file
                cells[1].text = ev.quote[:500]
        if r.caveats:
            doc.add_paragraph(f"Caveats: {r.caveats}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
